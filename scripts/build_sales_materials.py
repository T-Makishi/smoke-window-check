from __future__ import annotations

from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from reportlab.graphics.barcode.qr import QrCodeWidget
from reportlab.graphics.shapes import Drawing
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_PDF = ROOT / "output" / "pdf"
OUT_DOCX = ROOT / "output" / "docx"
HERO = ROOT / "assets" / "marketing" / "smoke-window-inspection-hero.png"
LP_DESKTOP = ROOT / "tmp" / "pdfs" / "screens" / "lp-desktop.png"
LP_MOBILE = ROOT / "tmp" / "pdfs" / "screens" / "lp-mobile.png"
TRIAL_URL = "https://smoke-window-check.pages.dev/trial.html"

GREEN = HexColor("#0E4B37")
GREEN_2 = HexColor("#176548")
INK = HexColor("#0C2940")
MUTED = HexColor("#60716B")
PAPER = HexColor("#F4F7F5")
LINE = HexColor("#D8E3DD")
GOLD = HexColor("#D4A748")
WHITE = colors.white

JAPANESE_FONT = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
pdfmetrics.registerFont(TTFont("JapaneseUI", JAPANESE_FONT))
FONT = "JapaneseUI"


def P(text, style, **kwargs):
    return Paragraph(text, style, **kwargs)


def styles():
    s = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title", parent=s["Title"], fontName=FONT, fontSize=25, leading=34, textColor=INK, alignment=TA_LEFT, spaceAfter=6),
        "cover_title": ParagraphStyle("cover_title", parent=s["Title"], fontName=FONT, fontSize=28, leading=38, textColor=WHITE, alignment=TA_LEFT, spaceAfter=7),
        "h1": ParagraphStyle("h1", parent=s["Heading1"], fontName=FONT, fontSize=18, leading=26, textColor=INK, spaceAfter=8),
        "h2": ParagraphStyle("h2", parent=s["Heading2"], fontName=FONT, fontSize=12.5, leading=18, textColor=GREEN, spaceBefore=5, spaceAfter=5),
        "body": ParagraphStyle("body", parent=s["BodyText"], fontName=FONT, fontSize=9.2, leading=15, textColor=INK, spaceAfter=5),
        "small": ParagraphStyle("small", parent=s["BodyText"], fontName=FONT, fontSize=7.5, leading=11.5, textColor=MUTED),
        "tiny": ParagraphStyle("tiny", parent=s["BodyText"], fontName=FONT, fontSize=6.6, leading=9.5, textColor=MUTED),
        "white": ParagraphStyle("white", parent=s["BodyText"], fontName=FONT, fontSize=9.5, leading=15, textColor=WHITE),
        "white_small": ParagraphStyle("white_small", parent=s["BodyText"], fontName=FONT, fontSize=7.5, leading=11, textColor=HexColor("#D5E4DE")),
        "kicker": ParagraphStyle("kicker", parent=s["BodyText"], fontName=FONT, fontSize=7.5, leading=10, textColor=GREEN, spaceAfter=7),
        "kicker_light": ParagraphStyle("kicker_light", parent=s["BodyText"], fontName=FONT, fontSize=7.5, leading=10, textColor=HexColor("#B2D7C8"), spaceAfter=7),
        "step_no": ParagraphStyle("step_no", parent=s["BodyText"], fontName=FONT, fontSize=8, leading=10, textColor=GREEN, alignment=TA_CENTER),
        "center": ParagraphStyle("center", parent=s["BodyText"], fontName=FONT, fontSize=8.5, leading=13, textColor=INK, alignment=TA_CENTER),
        "flyer_title": ParagraphStyle("flyer_title", parent=s["Title"], fontName=FONT, fontSize=26, leading=35, textColor=INK, alignment=TA_LEFT),
    }


S = styles()


class PageBand(Flowable):
    def __init__(self, height=7 * mm, color=GREEN):
        super().__init__(); self.width = 1; self.height = height; self.color = color

    def draw(self):
        self.canv.setFillColor(self.color); self.canv.rect(-18 * mm, 0, A4[0], self.height, fill=1, stroke=0)


class UiPanel(Flowable):
    def __init__(self, kind="customer", width=155 * mm, height=70 * mm):
        super().__init__(); self.width = width; self.height = height; self.kind = kind

    def draw(self):
        c = self.canv; w, h = self.width, self.height
        c.setFillColor(WHITE); c.setStrokeColor(LINE); c.roundRect(0, 0, w, h, 5, fill=1, stroke=1)
        c.setFillColor(INK); c.roundRect(0, h - 11 * mm, w, 11 * mm, 5, fill=1, stroke=0); c.rect(0, h - 11 * mm, w, 5, fill=1, stroke=0)
        c.setFont(FONT, 7.5); c.setFillColor(WHITE)
        c.drawString(7 * mm, h - 7 * mm, "排煙窓スマホ事前見積")
        if self.kind == "vendor":
            c.setFillColor(GREEN); c.roundRect(7 * mm, h - 22 * mm, 50 * mm, 7 * mm, 2, fill=1, stroke=0)
            c.setFillColor(WHITE); c.setFont(FONT, 6.8); c.drawString(11 * mm, h - 19.5 * mm, "試験利用中　残り7日")
            labels = ["会社名・受付メール", "料金・案内文", "ロゴ画像", "お客様への案内"]
            y = h - 31 * mm
            for i, label in enumerate(labels):
                c.setFillColor(PAPER); c.setStrokeColor(LINE); c.roundRect(7 * mm, y - 8 * mm, 65 * mm, 9 * mm, 2, fill=1, stroke=1)
                c.setFillColor(INK); c.setFont(FONT, 6.6); c.drawString(10 * mm, y - 4.5 * mm, label)
                y -= 11 * mm
            c.setFillColor(GREEN); c.roundRect(80 * mm, 13 * mm, 68 * mm, 32 * mm, 3, fill=1, stroke=0)
            c.setFillColor(WHITE); c.setFont(FONT, 8); c.drawString(86 * mm, 37 * mm, "お客様への案内を準備")
            c.setFont(FONT, 6.6); c.drawString(86 * mm, 30 * mm, "メール用リンクをコピー")
            c.drawString(86 * mm, 23 * mm, "LINE用QR画像を作成")
        else:
            c.setFillColor(INK); c.setFont(FONT, 7); c.drawCentredString(w / 2, h - 20 * mm, "写真と簡単な質問で")
            c.setFont(FONT, 15); c.drawCentredString(w / 2, h - 31 * mm, "修理費用の目安を確認できます")
            labels = ["入力は約3分", "事前チェック無料", "スマホで完結"]
            for i, label in enumerate(labels):
                x = 19 * mm + i * 42 * mm
                c.setFillColor(PAPER); c.roundRect(x, h - 46 * mm, 36 * mm, 10 * mm, 2, fill=1, stroke=0)
                c.setFillColor(GREEN); c.setFont(FONT, 6.5); c.drawCentredString(x + 18 * mm, h - 42 * mm, label)
            c.setFillColor(GREEN); c.roundRect(24 * mm, 8 * mm, w - 48 * mm, 9 * mm, 2, fill=1, stroke=0)
            c.setFillColor(WHITE); c.setFont(FONT, 7.5); c.drawCentredString(w / 2, 11.5 * mm, "事前チェックを始める  ›")


class Checklist(Flowable):
    def __init__(self, items, width=155 * mm):
        super().__init__(); self.items = items; self.width = width; self.height = len(items) * 9 * mm

    def draw(self):
        c = self.canv; y = self.height - 6 * mm
        for item in self.items:
            c.setFillColor(GREEN); c.circle(3 * mm, y + 1.5 * mm, 2.2 * mm, fill=1, stroke=0)
            c.setFillColor(WHITE); c.setFont(FONT, 5); c.drawCentredString(3 * mm, y, "✓")
            c.setFillColor(INK); c.setFont(FONT, 8.3); c.drawString(9 * mm, y, item)
            y -= 9 * mm


def manual_canvas(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(LINE); canvas.line(18 * mm, 16 * mm, A4[0] - 18 * mm, 16 * mm)
    canvas.setFont(FONT, 6.8); canvas.setFillColor(MUTED)
    canvas.drawString(18 * mm, 10.5 * mm, "PCSAPO / マキシ企画")
    canvas.drawRightString(A4[0] - 18 * mm, 10.5 * mm, f"{doc.title}  |  {doc.page}")
    canvas.restoreState()


def build_doc(path: Path, title: str):
    doc = BaseDocTemplate(str(path), pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=17 * mm, bottomMargin=21 * mm, title=title, author="PCSAPO / マキシ企画")
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates(PageTemplate(id="manual", frames=frame, onPage=manual_canvas))
    return doc


def cover(title, subtitle, audience, image_path=HERO):
    img = Image(str(image_path), width=174 * mm, height=98 * mm)
    title_block = Table([[P(audience, S["kicker_light"]), ""], [P(title, S["cover_title"]), ""], [P(subtitle, S["white"]), ""]], colWidths=[135 * mm, 39 * mm], rowHeights=[12 * mm, 34 * mm, 24 * mm])
    title_block.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), GREEN), ("LEFTPADDING", (0, 0), (-1, -1), 10 * mm), ("RIGHTPADDING", (0, 0), (-1, -1), 8 * mm), ("TOPPADDING", (0, 0), (-1, -1), 4 * mm), ("BOTTOMPADDING", (0, 0), (-1, -1), 2 * mm), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    return [img, title_block, Spacer(1, 12 * mm), P("このマニュアルでできること", S["h2"])]


def step_table(items):
    cells = []
    for idx, (title, body) in enumerate(items, 1):
        cells.append([P(f"{idx:02d}", S["step_no"]), P(f"<b>{escape(title)}</b><br/><font color='#60716B'>{escape(body)}</font>", S["body"])])
    t = Table(cells, colWidths=[14 * mm, 141 * mm], rowHeights=[None] * len(cells))
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (0, -1), PAPER), ("BOX", (0, 0), (-1, -1), .5, LINE), ("INNERGRID", (0, 0), (-1, -1), .5, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 4 * mm), ("RIGHTPADDING", (0, 0), (-1, -1), 4 * mm), ("TOPPADDING", (0, 0), (-1, -1), 3.5 * mm), ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5 * mm)]))
    return t


def note(text, color=PAPER):
    t = Table([[P(text, S["small"])]], colWidths=[155 * mm])
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), color), ("BOX", (0, 0), (-1, -1), .6, LINE), ("LEFTPADDING", (0, 0), (-1, -1), 5 * mm), ("RIGHTPADDING", (0, 0), (-1, -1), 5 * mm), ("TOPPADDING", (0, 0), (-1, -1), 3.5 * mm), ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5 * mm)]))
    return t


def build_vendor_pdf():
    path = OUT_PDF / "排煙窓スマホ事前見積_業者向け操作マニュアル.pdf"
    doc = build_doc(path, "業者向け操作マニュアル")
    story = cover("業者向け\n操作マニュアル", "無料体験の開始から、お客様用URL・QRコードの案内まで", "FOR VENDORS / 2026.08")
    story += [Checklist(["メール確認と初回パスコード登録", "会社情報・受付メール・料金・ロゴの設定", "お客様用URLとQRコードの案内", "試験期間の確認と本番利用申込"]), Spacer(1, 6 * mm), note("対象URLは会社ごとに発行されます。別会社のURLを使うと設定内容と受付先が異なるため、必ず自社専用URLを保管してください。"), PageBreak()]
    story += [P("01　無料体験を開始する", S["title"]), P("メール確認後に会社専用環境が発行されます。", S["body"]), Spacer(1, 4 * mm), step_table([("申込ページを開く", "7日間無料体験ページで会社名・担当者名・メール・電話番号を入力します。"), ("確認メールを開く", "30分以内に確認リンクを開きます。届かない場合は迷惑メールフォルダを確認します。"), ("業者設定URLを保存する", "画面とメールに表示された業者設定URLをブックマークします。"), ("初回パスコードを登録する", "登録メールで確認後、推測されにくい6〜8桁の数字を2回入力します。")]), Spacer(1, 8 * mm), UiPanel("vendor"), Spacer(1, 6 * mm), note("確認メールの送信元として Cloudflare の通知アドレスが表示される場合があります。確認コードを再送すると、古いコードは無効になります。"), PageBreak()]
    story += [P("02　会社情報と受付内容を設定する", S["title"]), P("設定を保存した時点の内容が、会社専用のお客様用URLへ反映されます。", S["body"]), Spacer(1, 4 * mm), step_table([("基本情報", "会社名、電話番号、所在地など、お客様へ表示する情報を入力します。"), ("問診受付メール", "お客様から問診を受け取る実際の業務用メールを設定します。"), ("料金・訪問条件", "現地調査費、概算料金、キャンセル・不在時の説明を自社条件に合わせます。"), ("ロゴ画像", "PNG・JPEG・WebP、1MB以下を選択します。共有用に自動圧縮されます。"), ("設定を保存", "画面下の保存ボタンを押し、完了表示を確認します。")]), Spacer(1, 7 * mm), P("設定後の確認項目", S["h2"]), Checklist(["会社名・ロゴ・連絡先が正しい", "受付メールが実際に確認できるアドレス", "現地調査費と概算料金が最新", "注意事項が自社の運用と一致"]), Spacer(1, 3 * mm), note("受付メールを変更しただけでは、以前お客様へ案内したURLの設定は変わりません。設定保存後に表示される会社専用URLを使用してください。"), PageBreak()]
    story += [P("03　お客様へ案内・問診を受信する", S["title"]), step_table([("メールで案内", "「設定を保存してメール用リンクをコピー」を押し、案内文へ貼り付けます。表示は短いリンクになります。"), ("LINE・対面で案内", "「LINE用QR画像を作成」を押し、表示されたQRを見せるかPNG画像として送ります。"), ("事前にテスト", "自分のスマートフォンでお客様用URLを開き、会社表示と受付メールを確認します。"), ("問診を受信", "お客様がメール作成画面を開くと受付先が自動入力されます。写真・動画はお客様が添付します。")]), Spacer(1, 7 * mm), P("試験期間と本番利用", S["h2"]), P("試験版では画面上部に残り日数が表示されます。期限後はお客様用画面が案内表示へ切り替わります。本番利用を希望する場合は、業者設定内の「本番利用を申し込む」から担当者名と連絡事項を送信してください。送信だけで契約や課金は確定しません。", S["body"]), Spacer(1, 4 * mm), note("パスコードを忘れた場合：業者ログイン画面の「パスコードを忘れた場合」から再設定用URLをコピーしてログアウトし、登録メールで再認証後、新しい6〜8桁を登録します。"), Spacer(1, 7 * mm), P("運用前チェック", S["h2"]), Checklist(["PC・タブレット・スマートフォンで表示確認", "自社宛にテスト問診を1件送信", "写真・動画の添付手順を担当者間で共有", "専用URLとQR画像を安全な場所へ保管"])]
    doc.build(story)
    return path


def build_customer_pdf():
    path = OUT_PDF / "排煙窓スマホ事前見積_ご依頼者向け操作マニュアル.pdf"
    doc = build_doc(path, "ご依頼者向け操作マニュアル")
    story = cover("ご依頼者向け\n操作マニュアル", "写真と簡単な質問で、排煙窓の状況を修理会社へ伝える方法", "FOR CUSTOMERS / 2026.08")
    story += [Checklist(["ログイン・アプリのインストールは不要", "入力時間の目安は約3分", "写真・動画は送信時に自分で添付", "途中の相談内容は利用中の端末へ保存"]), Spacer(1, 6 * mm), note("このアプリで表示する金額は概算です。正式な見積りには、修理会社による現地確認が必要になる場合があります。"), PageBreak()]
    story += [P("01　事前チェックを進める", S["title"]), UiPanel("customer"), Spacer(1, 7 * mm), step_table([("案内URLまたはQRを開く", "修理会社から届いた専用URLを開きます。会社名が正しいことを確認します。"), ("お客様情報を入力", "氏名・会社名・施設名・住所・連絡先など、表示された項目を入力します。"), ("症状を複数選ぶ", "動きが重い、開かない、閉まらないなど、当てはまる症状をすべて選びます。"), ("設置状況を入力", "台数、高さ、窓の種類、緊急度など、分かる範囲で回答します。"), ("写真・動画を選ぶ", "全体、操作部、故障箇所が分かる画像を選びます。危険な場所では撮影しないでください。")]), PageBreak()]
    story += [P("02　内容を保存・送信する", S["title"]), step_table([("概算結果を確認", "料金の目安、現地調査費、注意事項を確認します。正式金額ではありません。"), ("端末へ保存", "後で続ける場合は「この端末に保存」を使用します。写真・動画の実データは保存されません。"), ("印刷する", "印刷ボタンを押します。開かない場合はブラウザのメニューから「印刷」または「プリント」を選びます。"), ("メールを作成", "送信ボタンでメール画面を開きます。受付先と問診内容が入力されていることを確認します。"), ("写真・動画を添付", "メール作成画面のクリップまたは添付ボタンから、選択した写真・動画を追加して送信します。")]), Spacer(1, 7 * mm), P("LINEで送る場合", S["h2"]), P("端末の共有機能に対応している場合は写真・動画をLINEへ共有できます。問診本文と送信先が自動で完全に引き継がれない場合があるため、修理会社から指定された送信方法を優先してください。", S["body"]), Spacer(1, 4 * mm), note("個人情報を含みます。共用端末では保存しないでください。ブラウザの閲覧データを削除すると保存した相談も消えます。送信前に宛先が依頼する修理会社のメールであることを確認してください。"), Spacer(1, 7 * mm), P("撮影時の安全", S["h2"]), Checklist(["脚立や高所へ無理に上がらない", "排煙窓を強く操作しない", "煙・火災時は撮影より避難と通報を優先", "不明な場合は修理会社へ直接連絡"])]
    doc.build(story)
    return path


def qr_drawing(data, size=34 * mm):
    qr = QrCodeWidget(data)
    b = qr.getBounds(); w, h = b[2] - b[0], b[3] - b[1]
    d = Drawing(size, size, transform=[size / w, 0, 0, size / h, 0, 0]); d.add(qr); return d


def build_flyer_pdf():
    path = OUT_PDF / "排煙窓スマホ事前見積_A4営業チラシ.pdf"
    doc = build_doc(path, "A4営業チラシ")
    hero = Image(str(HERO), width=174 * mm, height=78 * mm)
    headline = Table([[P("排煙窓の事前確認を、<br/><font color='#0E4B37'>スマートフォンで。</font>", S["flyer_title"]), P("現地へ行く前に、症状・写真・動画・連絡先をまとめて受け取る。排煙窓の修理・点検会社向けWebアプリです。", S["body"])]], colWidths=[108 * mm, 58 * mm])
    headline.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "BOTTOM"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (0, 0), 8 * mm), ("RIGHTPADDING", (1, 0), (1, 0), 0), ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
    benefits = []
    for n, title, body in [("01", "聞き取りを標準化", "複数症状・設置状況・連絡先を同じ順序で確認。"), ("02", "自社仕様に変更", "受付メール・ロゴ・料金・案内文を会社別に設定。"), ("03", "メール・LINEで案内", "短い表示リンクとQR画像で、お客様へ簡単に共有。")]:
        benefits.append([P(n, S["step_no"]), P(f"<b>{title}</b><br/><font color='#60716B'>{body}</font>", S["small"])])
    benefit_table = Table(benefits, colWidths=[13 * mm, 42 * mm] * 0 + [13 * mm, 42 * mm], repeatRows=0)
    # Arrange benefits horizontally using nested two-column cards.
    cards = []
    for n, title, body in [("01", "聞き取りを標準化", "複数症状・設置状況・連絡先を同じ順序で確認。"), ("02", "自社仕様に変更", "受付メール・ロゴ・料金・案内文を会社別に設定。"), ("03", "メール・LINEで案内", "短い表示リンクとQR画像で、お客様へ簡単に共有。")]:
        cards.append([P(n, S["step_no"]), P(f"<b>{title}</b><br/><font color='#60716B'>{body}</font>", S["small"])])
    card_tables = []
    for card in cards:
        t = Table([card], colWidths=[10 * mm, 42 * mm], rowHeights=[27 * mm])
        t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), PAPER), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 3 * mm), ("RIGHTPADDING", (0, 0), (-1, -1), 3 * mm), ("BOX", (0, 0), (-1, -1), .5, LINE)])); card_tables.append(t)
    row = Table([card_tables], colWidths=[54 * mm, 54 * mm, 54 * mm])
    row.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 2 * mm), ("RIGHTPADDING", (-1, 0), (-1, -1), 0)]))
    cta = Table([[P("7日間無料体験", ParagraphStyle("cta1", parent=S["h1"], textColor=WHITE, fontSize=17, leading=22)), P("クレジットカード不要<br/>自動課金なし", S["white_small"]), qr_drawing(TRIAL_URL, 31 * mm)]], colWidths=[65 * mm, 55 * mm, 38 * mm], rowHeights=[38 * mm])
    cta.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), GREEN), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 6 * mm), ("RIGHTPADDING", (0, 0), (-1, -1), 3 * mm), ("TOPPADDING", (0, 0), (-1, -1), 3 * mm), ("BOTTOMPADDING", (0, 0), (-1, -1), 3 * mm)]))
    story = [P("排煙窓の修理・点検会社向け", S["kicker"]), headline, Spacer(1, 7 * mm), hero, Spacer(1, 7 * mm), row, Spacer(1, 7 * mm), P("お客様はログイン不要。URLを開いて約3分で入力。", S["h1"]), P("問診内容はお客様の端末へ保存し、写真・動画はメール送信時に添付します。運営側へ顧客の問診データを蓄積しない、試験運用向けのシンプルな構成です。", S["body"]), Spacer(1, 6 * mm), cta, Spacer(1, 4 * mm), P(f"無料体験：{TRIAL_URL}", S["tiny"])]
    doc.build(story)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def set_doc_font(run, size=10, bold=False, color="0C2940"):
    run.font.name = "Arial Unicode MS"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)


def add_doc_title(doc, title, subtitle):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(5)
    set_doc_font(p.add_run(title), 26, True, "0C2940")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(18); set_doc_font(p.add_run(subtitle), 11, False, "60716B")


def add_doc_heading(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6); set_doc_font(p.add_run(text), 17, True, "0E4B37"); return p


def add_doc_text(doc, text, bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.22; set_doc_font(p.add_run(text), 10, bold); return p


def add_doc_steps(doc, items):
    table = doc.add_table(rows=0, cols=2); table.autofit = False
    for idx, (title, body) in enumerate(items, 1):
        cells = table.add_row().cells; cells[0].width = Mm(15); cells[1].width = Mm(145)
        set_cell_shading(cells[0], "EAF2EE"); cells[0].vertical_alignment = cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_doc_font(p.add_run(f"{idx:02d}"), 9, True, "0E4B37")
        p = cells[1].paragraphs[0]; set_doc_font(p.add_run(title + "\n"), 10, True); set_doc_font(p.add_run(body), 9, False, "60716B")
    return table


def setup_docx(title):
    doc = Document(); section = doc.sections[0]; section.page_width = Mm(210); section.page_height = Mm(297); section.top_margin = Mm(16); section.bottom_margin = Mm(18); section.left_margin = Mm(18); section.right_margin = Mm(18)
    section.header_distance = Mm(8); section.footer_distance = Mm(8)
    hp = section.header.paragraphs[0]; set_doc_font(hp.add_run(title), 8, True, "60716B")
    fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_doc_font(fp.add_run("PCSAPO / マキシ企画"), 8, False, "60716B")
    doc.core_properties.title = title; doc.core_properties.author = "PCSAPO / マキシ企画"
    return doc


def build_docx_manuals():
    vendor = setup_docx("排煙窓スマホ事前見積｜業者向け操作マニュアル")
    add_doc_title(vendor, "業者向け操作マニュアル", "無料体験の開始から、お客様用URL・QRコードの案内まで")
    vendor.add_picture(str(HERO), width=Mm(174)); add_doc_text(vendor, "対象：排煙窓の修理・点検会社　｜　7日間無料体験・本番利用", True)
    vendor.add_page_break(); add_doc_heading(vendor, "1. 無料体験を開始する")
    add_doc_steps(vendor, [("申込フォームを送信", "会社名、担当者、メール、電話番号を入力します。"), ("メール確認", "30分以内に確認リンクを開きます。"), ("業者設定URLを保存", "会社専用URLをブックマークします。"), ("パスコード登録", "6〜8桁の推測されにくい数字を登録します。")])
    add_doc_heading(vendor, "2. 会社情報を設定する"); add_doc_steps(vendor, [("会社情報・受付メール", "顧客表示と問診の受付先を入力します。"), ("料金・訪問条件", "現地調査費、概算料金、注意事項を更新します。"), ("ロゴ", "PNG・JPEG・WebP、1MB以下を選択します。"), ("設定保存", "保存完了を確認してから案内URLを発行します。")])
    vendor.add_page_break(); add_doc_heading(vendor, "3. お客様へ案内する")
    add_doc_steps(vendor, [("メール", "メール用リンクをコピーして案内文へ貼り付けます。"), ("LINE・対面", "QR画像を作成して表示またはPNG送信します。"), ("事前テスト", "自分の端末から会社表示と送信先を確認します。"), ("問診受信", "お客様がメールへ写真・動画を添付して送信します。")])
    add_doc_heading(vendor, "4. 試験期間・本番利用・パスコード"); add_doc_text(vendor, "試験版では上部に残り日数を表示します。期限後は案内画面へ切り替わります。本番利用は業者設定から申し込み、運営者が条件確認後に切り替えます。パスコードを忘れた場合は、登録メールで再認証して再設定します。")
    vendor_path = OUT_DOCX / "排煙窓スマホ事前見積_業者向け操作マニュアル.docx"; vendor.save(vendor_path)

    customer = setup_docx("排煙窓スマホ事前見積｜ご依頼者向け操作マニュアル")
    add_doc_title(customer, "ご依頼者向け操作マニュアル", "写真と簡単な質問で、排煙窓の状況を修理会社へ伝える方法")
    customer.add_picture(str(HERO), width=Mm(174)); add_doc_text(customer, "ログイン不要・アプリのインストール不要・入力目安約3分", True)
    customer.add_page_break(); add_doc_heading(customer, "1. 事前チェックを進める")
    add_doc_steps(customer, [("URL・QRを開く", "案内した修理会社名が表示されていることを確認します。"), ("お客様情報", "連絡先や施設情報を入力します。"), ("症状を選ぶ", "当てはまる症状を複数選択できます。"), ("設置状況", "台数、高さ、窓の種類など分かる範囲で回答します。"), ("写真・動画", "危険のない範囲で全体・操作部・故障箇所を撮影します。")])
    add_doc_heading(customer, "2. 内容を送信する"); add_doc_steps(customer, [("概算確認", "金額は目安です。正式見積りには現地確認が必要な場合があります。"), ("保存・印刷", "端末保存またはブラウザの印刷機能を使用します。"), ("メール作成", "受付先と問診本文を確認します。"), ("添付して送信", "写真・動画をメールへ追加して送信します。")])
    add_doc_heading(customer, "安全と個人情報"); add_doc_text(customer, "高所へ無理に上がらず、排煙窓を強く操作しないでください。共用端末には相談内容を保存しないでください。送信前に宛先が依頼する修理会社のメールアドレスであることを確認します。")
    customer_path = OUT_DOCX / "排煙窓スマホ事前見積_ご依頼者向け操作マニュアル.docx"; customer.save(customer_path)
    return vendor_path, customer_path


def main():
    OUT_PDF.mkdir(parents=True, exist_ok=True); OUT_DOCX.mkdir(parents=True, exist_ok=True)
    paths = [build_vendor_pdf(), build_customer_pdf(), build_flyer_pdf()]
    for p in paths: print(p)


if __name__ == "__main__":
    main()
